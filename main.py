from PySide6 import QtWidgets, QtCore, QtGui
import sys
from pathlib import Path
from docx import Document
import textwrap
from pynput import keyboard
import pyperclip
import pyautogui as pag


class WarningBox(QtWidgets.QMessageBox):
    def __init__(self, title, text, ok=False):
        super().__init__()
        self.setWindowTitle(title)
        longest = max(map(len, text.split()))
        longest = longest if longest > 50 else 50
        wrapped_text = textwrap.wrap(text, width=longest)
        width = 200 + longest * 4
        height = 50 + 20 * (len(wrapped_text) - 1)
        self.setText("\n".join(wrapped_text))
        self.setIcon(QtWidgets.QMessageBox.Icon.Warning)
        self.setStandardButtons(QtWidgets.QMessageBox.StandardButton.Ok)
        self.setStyleSheet("""
                    QMessageBox QLabel#qt_msgbox_label {
                        min-width: %spx;
                        min-height: %spx;
                        word-wrap: break-word;
                    }
                """ % (width, height))
        if ok:
            self.setStandardButtons(QtWidgets.QMessageBox.StandardButton.Ok |
                                    QtWidgets.QMessageBox.StandardButton.Cancel)
        self.exec()


class Window(QtWidgets.QMainWindow):
    hotkey_triggered = QtCore.Signal()

    def __init__(self):
        super().__init__()
        self.setWindowTitle("Viber Copier")
        self.setFixedSize(800, 600)
        self.central = QtWidgets.QWidget()
        self.setCentralWidget(self.central)
        self.layout = QtWidgets.QVBoxLayout(self.central)

        self.font = QtGui.QFont("Arial", 14)
        self.file_path = None
        self.file_path_copy = None
        self.last_text = ""
        self.listener = keyboard.Listener(on_press=self.on_press)
        self.listening = False

        self.hotkey_triggered.connect(self.auto_paste)

        self.open_file_button = QtWidgets.QPushButton("Open")
        self.open_file_button.setFont(self.font)
        self.open_file_button.setMinimumWidth(200)
        self.open_file_button.clicked.connect(self.open_file)
        self.layout.addWidget(self.open_file_button, alignment=QtCore.Qt.AlignmentFlag.AlignLeft)

        self.key_sniffer = QtWidgets.QPushButton("Not listening")
        self.key_sniffer.clicked.connect(self.listen)
        self.layout.addWidget(self.key_sniffer)

        self.file_path_label = QtWidgets.QLabel("No File Selected")
        self.file_path_label.setFont(self.font)
        self.layout.addWidget(self.file_path_label)

        self.text_editor = QtWidgets.QTextEdit()
        self.text_editor.setFont(self.font)
        self.layout.addWidget(self.text_editor)

        self.paste_button = QtWidgets.QPushButton("Paste")
        self.paste_button.setFont(self.font)
        self.paste_button.setMinimumWidth(200)
        self.paste_button.clicked.connect(self.paste_text)
        self.layout.addWidget(self.paste_button, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)

    def on_press(self, key):
        try:
            if key.char == "q":
                self.hotkey_triggered.emit()
        except AttributeError:
            pass

    def listen(self):
        if not self.listening:
            self.listening = True
            self.key_sniffer.setText("Listening")
            self.listener.start()
        else:
            self.listening = False
            self.key_sniffer.setText("Not listening")
            self.listener.stop()

    def open_file(self):
        name, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Open File", "", "DOCX (*.docx);;")
        name_path = Path(name)
        if name and name_path.exists():
            self.file_path = name_path
            self.file_path_copy = Path(self.file_path).parent / f"{self.file_path.stem}_copy{self.file_path.suffix}"
            drive = self.file_path.drive
            parent = self.file_path.parent.name
            filename = self.file_path.name
            self.file_path_label.setText(f"{drive}\\...\\{parent}\\{filename}")

    @QtCore.Slot()
    def auto_paste(self):
        pag.hotkey("ctrl", "c")
        QtCore.QTimer.singleShot(500, self.auto_paste_next)

    def auto_paste_next(self):
        self.text_editor.setText(pyperclip.paste())
        QtCore.QTimer.singleShot(500, self.paste_text)

    def paste_text(self):
        text = self.text_editor.toPlainText()
        if self.file_path is None:
            WarningBox("Error", "No file selected")
            return
        if not text:
            WarningBox("Error", "No text to paste")
            return
        if text == self.last_text:
            msg_box = WarningBox("Warning", "You are about to insert the same block of text", True)
            clicked = msg_box.clickedButton()
            if clicked == msg_box.button(QtWidgets.QMessageBox.StandardButton.Ok):
                pass
            elif clicked == msg_box.button(QtWidgets.QMessageBox.StandardButton.Cancel):
                return
            else:
                return
        self.paste_button.setEnabled(False)
        try:
            doc = Document(str(self.file_path))
            if doc.paragraphs:
                new_paragraph = doc.paragraphs[0].insert_paragraph_before()
            else:
                new_paragraph = doc.add_paragraph()
            new_paragraph.text = text
            doc.save(str(self.file_path_copy))
            name = self.file_path.name
            self.file_path.unlink()
            self.file_path_copy.rename(name)
            self.text_editor.clear()
            self.last_text = text
        except Exception as e:
            WarningBox(f"{type(e).__name__}", f"{e}")

        self.paste_button.setEnabled(True)

    def closeEvent(self, event):
        if self.listener.running:
            self.listener.stop()
        event.accept()


if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    window = Window()
    window.show()
    sys.exit(app.exec())
